"""
DOCX/DOC parser (python-docx for .docx; .doc is registered as an alias
that raises a clear UnsupportedFileTypeError-adjacent message, since
legacy binary .doc requires a different library than python-docx
supports and is explicitly out of scope for this sprint's dependency
set — see class docstring below).
"""

from datetime import datetime, timezone

from ingestion.base import BaseParser, ParserRegistry
from ingestion.exceptions import CorruptedFileError
from ingestion.result import DocumentMetadata, OcrInfo, ParsedDocument, ProcessingInfo
from ingestion.utils import count_characters, count_words, detect_language


def _extract_tables(document) -> list[list[list[str]]]:
    tables = []
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)
    return tables


def _extract_headers_and_footers(document) -> tuple[list[str], list[str]]:
    headers, footers = [], []
    for section in document.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
        footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
        if header_text:
            headers.append(header_text)
        if footer_text:
            footers.append(footer_text)
    return headers, footers


@ParserRegistry.register
class DocxParser(BaseParser):
    extension = "docx"

    def parse(self, file_path: str, document_id: str) -> ParsedDocument:
        started_at = datetime.now(timezone.utc)
        warnings: list[str] = []

        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = docx.Document(file_path)
        except PackageNotFoundError as exc:
            raise CorruptedFileError(
                f"'{file_path}' is not a valid DOCX file: {exc}"
            ) from exc

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        headers, footers = _extract_headers_and_footers(document)
        tables = _extract_tables(document)

        text_parts = list(paragraphs)
        if headers:
            text_parts = headers + text_parts
        if footers:
            text_parts = text_parts + footers
        text = "\n".join(text_parts)

        core_properties = document.core_properties
        metadata = DocumentMetadata(
            title=core_properties.title or None,
            author=core_properties.author or None,
            created_at=core_properties.created.isoformat() if core_properties.created else None,
            modified_at=core_properties.modified.isoformat() if core_properties.modified else None,
            language=detect_language(text),
            page_count=None,  # python-docx cannot determine rendered page count.
            extension=self.extension,
            parser_used=self.__class__.__name__,
            ocr_used=False,
            character_count=count_characters(text),
            word_count=count_words(text),
        )

        if metadata.page_count is None:
            warnings.append(
                "Page count is not available for DOCX files — python-docx has no "
                "concept of rendered pages without a layout engine."
            )

        finished_at = datetime.now(timezone.utc)
        duration = (finished_at - started_at).total_seconds()
        metadata.parsing_time_seconds = duration

        return ParsedDocument(
            document_id=document_id,
            text=text,
            metadata=metadata,
            tables=tables,
            images=[],
            ocr=OcrInfo(used=False),
            processing=ProcessingInfo(
                started_at=started_at,
                finished_at=finished_at,
                duration_seconds=duration,
                warnings=warnings,
            ),
        )


@ParserRegistry.register
class DocParser(BaseParser):
    """
    Legacy binary .doc format.

    python-docx only supports the modern .docx (OOXML) format; parsing
    legacy binary .doc requires a different dependency (e.g.
    antiword/textract/LibreOffice) that is not part of this sprint's
    approved dependency set. Rather than silently mis-parsing or
    pretending .doc works, this parser raises a clear, actionable
    CorruptedFileError-style message identifying the real limitation.
    """

    extension = "doc"

    def parse(self, file_path: str, document_id: str) -> ParsedDocument:
        raise CorruptedFileError(
            "Legacy .doc (binary Word 97-2003) format is not supported by the "
            "current parser dependency set (python-docx only supports .docx). "
            "Please convert the file to .docx and re-upload."
        )
